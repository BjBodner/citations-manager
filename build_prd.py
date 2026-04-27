"""Generate a plain Hebrew RTL PRD document for the citations manager."""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


FONT = "Arial"
BODY_PT = 11
H1_PT = 18
H2_PT = 14
H3_PT = 12


def set_rtl_paragraph(paragraph):
    pPr = paragraph._element.get_or_add_pPr()
    bidi = OxmlElement("w:bidi")
    pPr.append(bidi)


def style_run(run, size_pt=BODY_PT, bold=False, color=None):
    run.font.name = FONT
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color
    rPr = run._element.get_or_add_rPr()
    rtl = OxmlElement("w:rtl")
    rPr.append(rtl)
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:cs"), FONT)
    rFonts.set(qn("w:ascii"), FONT)
    rFonts.set(qn("w:hAnsi"), FONT)
    rPr.append(rFonts)
    szCs = OxmlElement("w:szCs")
    szCs.set(qn("w:val"), str(size_pt * 2))
    rPr.append(szCs)
    bCs = OxmlElement("w:bCs")
    if bold:
        rPr.append(bCs)


def add_heading(doc, text, level=1):
    sizes = {1: H1_PT, 2: H2_PT, 3: H3_PT}
    p = doc.add_paragraph()
    set_rtl_paragraph(p)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(text)
    style_run(run, size_pt=sizes[level], bold=True)
    p.paragraph_format.space_before = Pt(12 if level > 1 else 18)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_para(doc, text, bold=False):
    p = doc.add_paragraph()
    set_rtl_paragraph(p)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(text)
    style_run(run, bold=bold)
    p.paragraph_format.space_after = Pt(4)
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    set_rtl_paragraph(p)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.left_indent = Cm(0)
    p.paragraph_format.right_indent = Cm(0.6 + 0.6 * level)
    run = p.add_run(text)
    style_run(run)
    return p


def add_bullet_with_bold_label(doc, label, rest, level=0):
    p = doc.add_paragraph(style="List Bullet")
    set_rtl_paragraph(p)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.right_indent = Cm(0.6 + 0.6 * level)
    r1 = p.add_run(label)
    style_run(r1, bold=True)
    r2 = p.add_run(rest)
    style_run(r2)
    return p


def set_table_rtl(table):
    tblPr = table._element.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        table._element.insert(0, tblPr)
    bidiVisual = OxmlElement("w:bidiVisual")
    tblPr.append(bidiVisual)


def fill_cell(cell, text, bold=False, header=False):
    cell.text = ""
    p = cell.paragraphs[0]
    set_rtl_paragraph(p)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(text)
    style_run(run, bold=bold or header)
    if header:
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "D9E2F3")
        tcPr.append(shd)


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    set_table_rtl(table)
    for i, h in enumerate(headers):
        fill_cell(table.rows[0].cells[i], h, header=True)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            fill_cell(table.rows[1 + ri].cells[ci], val)
    return table


def main():
    doc = Document()

    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)

    sectPr = section._sectPr
    bidi = OxmlElement("w:bidi")
    sectPr.append(bidi)

    # Title
    title_p = doc.add_paragraph()
    set_rtl_paragraph(title_p)
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run("מסמך אפיון – מערכת ניהול ציטוטים אישית")
    style_run(title_run, size_pt=22, bold=True)
    title_p.paragraph_format.space_after = Pt(6)

    sub_p = doc.add_paragraph()
    set_rtl_paragraph(sub_p)
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub_p.add_run("Citations Manager – PRD")
    style_run(sub_run, size_pt=12, bold=False, color=RGBColor(0x55, 0x55, 0x55))
    sub_p.paragraph_format.space_after = Pt(4)

    meta_p = doc.add_paragraph()
    set_rtl_paragraph(meta_p)
    meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_run = meta_p.add_run("גרסה 1.0 · 27 באפריל 2026 · שימוש אישי")
    style_run(meta_run, size_pt=10, color=RGBColor(0x77, 0x77, 0x77))
    meta_p.paragraph_format.space_after = Pt(18)

    # 1. Executive summary
    add_heading(doc, "1. תקציר מנהלים", level=1)
    add_para(
        doc,
        "מסמך זה מאפיין מערכת אישית לניהול ציטוטים (Prior Art) עבור בוחן פטנטים ברשות הפטנטים. "
        "המערכת היא אפליקציית web מקומית הרצה בדפדפן, עם לוח Kanban ויזואלי, "
        "תמיכה בריבוי תיקים (בקשות פטנט), ואלמנטי משחוק (gamification) שמטרתם לשפר מוטיבציה ועמידה ביעדים יומיים.",
    )
    add_bullet_with_bold_label(doc, "מטרה: ", "כלי אישי כיפי, מהיר וגמיש לסיווג ציטוטים בעבודת בחינת פטנטים.")
    add_bullet_with_bold_label(doc, "קהל יעד: ", "המשתמש בלבד (שימוש אישי, לא מוצר).")
    add_bullet_with_bold_label(doc, "סוג המוצר: ", "אפליקציית React מקומית הרצה בדפדפן, ללא שרת.")

    # 2. Background
    add_heading(doc, "2. רקע ובעיה", level=1)
    add_para(
        doc,
        "בעבודת בחינת בקשות פטנט, הבוחן עובר על עשרות עד מאות ציטוטים של פרסומים מוקדמים (Prior Art) "
        "ומסווג כל אחד מהם כרלוונטי או לא רלוונטי. בתוך הציטוטים הרלוונטיים יש לסווג לפי הסיווגים הבינלאומיים: "
        "X (פוסל אדום), Y (פוסל בשילוב), I (כחול / מעניין) ו-A (רקע כללי).",
    )
    add_para(doc, "הבעיות בתהליך הקיים:")
    add_bullet(doc, "ניהול הציטוטים מתבצע ידנית במספר קבצים מפוזרים, מה שמקשה על מעקב.")
    add_bullet(doc, "אין תצוגה ויזואלית של מצב התקדמות.")
    add_bullet(doc, "אין מנגנון מוטיבציה שמעודד עמידה ביעדים יומיים.")
    add_bullet(doc, "אין מקום מרכזי שמפריד בין תיקים שונים שנבחנים במקביל.")

    # 3. Goals
    add_heading(doc, "3. מטרות והצלחה", level=1)
    add_bullet_with_bold_label(doc, "יעד פונקציונלי: ", "ניהול 100% מהציטוטים בלוח אחד לכל תיק, ללא קבצים חיצוניים.")
    add_bullet_with_bold_label(doc, "יעד חוויתי: ", "הגדלת מוטיבציה יומית – עמידה ביעד יומי ושמירה על Streak.")
    add_bullet_with_bold_label(doc, "יעד טכני: ", "אפליקציה שעובדת לחלוטין מקומית (localStorage), ללא תלות בשרת או באינטרנט.")
    add_bullet_with_bold_label(doc, "יעד פרטיות: ", "כל הנתונים נשמרים אך ורק במחשב המשתמש, ללא שליחה החוצה.")

    # 4. Persona
    add_heading(doc, "4. משתמש יעד (Persona)", level=1)
    add_para(
        doc,
        "בוחן פטנטים יחיד (המשתמש), בעל ידע מקצועי בתחום הפטנטים, "
        "המעוניין בכלי פרודוקטיביות אישי. המשתמש עובד מול מספר תיקים במקביל ומעוניין בהפרדה ברורה ביניהם.",
    )

    # 5. Functional requirements
    add_heading(doc, "5. דרישות פונקציונליות", level=1)

    add_heading(doc, "5.1 לוח Kanban", level=2)
    add_bullet(doc, "הלוח כולל 4 עמודות לפי הסטטוס: \u201cטרם התחיל\u201d, \u201cבביצוע\u201d, \u201cרלוונטי\u201d, \u201cלא רלוונטי\u201d.")
    add_bullet(doc, "בעמודת \u201cרלוונטי\u201d מופיעים Tabs בראש העמודה לסיווג: X / Y / I / A. בכל פעם נראים הציטוטים של ה-Tab הפעיל.")
    add_bullet(doc, "ניתן לגרור כרטיסים בין עמודות ובין Tabs באמצעות Drag & Drop.")
    add_bullet(doc, "כל עמודה מציגה מונה מספר הכרטיסים שבה.")

    add_heading(doc, "5.2 כרטיס ציטוט", level=2)
    add_para(doc, "השדות בכרטיס:")
    add_bullet(doc, "מספר פרסום (Publication Number) – לדוגמה US1234567B2.")
    add_bullet(doc, "לינק (URL) לציטוט המקורי.")
    add_bullet(doc, "כותרת (Title).")
    add_bullet(doc, "תקציר (Abstract).")
    add_bullet(doc, "הערות אישיות (טקסט חופשי).")
    add_bullet(doc, "סטטוס – העמודה בה הכרטיס נמצא.")
    add_bullet(doc, "תת-סיווג (X / Y / I / A) – רלוונטי רק כאשר הסטטוס הוא \u201cרלוונטי\u201d.")
    add_bullet(doc, "חותמות זמן: created, updated, completed.")

    add_heading(doc, "5.3 ריבוי תיקים (בקשות פטנט)", level=2)
    add_bullet(doc, "כל תיק = לוח Kanban נפרד ומבודד הקשור לבקשת פטנט אחת שהבוחן בודק.")
    add_bullet(doc, "מסך ראשי (\u201cWorkspace Switcher\u201d) מציג את רשימת התיקים, ומאפשר ליצור תיק חדש, לבחור תיק קיים, לשנות שם או למחוק.")
    add_bullet(doc, "בכל תיק יש שדה חיפוש/פילטר פנימי לציטוטים שלו.")
    add_bullet(doc, "ניווט מהיר בין תיקים מבלי לאבד את מצב הלוח.")

    add_heading(doc, "5.4 הוספת ציטוטים", level=2)
    add_heading(doc, "5.4.1 הוספה ידנית", level=3)
    add_para(doc, "כפתור \u201c+ ציטוט חדש\u201d פותח טופס עם כל השדות שבסעיף 5.2. בעת השמירה, הציטוט נכנס לעמודת \u201cטרם התחיל\u201d של התיק הפעיל.")

    add_heading(doc, "5.4.2 ייבוא CSV", level=3)
    add_bullet(doc, "פורמט קשיח עם כותרות מוגדרות (Header Row).")
    add_bullet(doc, "הכותרות: publication_number, title, abstract, link, notes.")
    add_bullet(doc, "תצוגה מקדימה של השורות לפני אישור הייבוא.")
    add_bullet(doc, "ולידציה: שורות שגויות נפסלות וההודעה למשתמש מציינת איזה שורות וטעות.")
    add_bullet(doc, "כפתור \u201cהורד דוגמת CSV\u201d ישירות מהמסך, עם 2-3 שורות לדוגמה.")
    add_bullet(doc, "כל הציטוטים שיובאו נכנסים לעמודת \u201cטרם התחיל\u201d של התיק הפעיל.")

    add_heading(doc, "5.5 פתיחת לינק (התנהגות אוטומטית)", level=2)
    add_bullet(doc, "לחיצה על הלינק בכרטיס שנמצא בעמודת \u201cטרם התחיל\u201d \u2192 הכרטיס עובר אוטומטית ל\u201cבביצוע\u201d והלינק נפתח בטאב חדש.")
    add_bullet(doc, "לחיצה על לינק בכרטיס בעמודה אחרת \u2192 רק פותחת את הלינק בטאב חדש, ללא שינוי סטטוס.")
    add_bullet(doc, "המעבר מסומן ויזואלית (אנימציה קצרה) כדי שהמשתמש יראה את ההעברה.")

    add_heading(doc, "5.6 ייצוא וייבוא JSON (גיבוי קריטי)", level=2)
    add_bullet(doc, "כפתור \u201cייצוא גיבוי\u201d מוריד קובץ JSON מלא של כל הלוחות, התיקים והגדרות המשתמש.")
    add_bullet(doc, "כפתור \u201cייבוא גיבוי\u201d טוען קובץ JSON. למשתמש תינתן בחירה: החלפה מלאה של הנתונים הקיימים, או מיזוג.")
    add_bullet(doc, "גיבוי אוטומטי יומי – המערכת שומרת snapshot יומי ב-localStorage עם תאריך, ושומרת עד 7 ימים אחרונים.")
    add_bullet(doc, "אזהרה למשתמש כשה-localStorage מתקרב לקיבולת המקסימלית.")

    add_heading(doc, "5.7 מוטיבציה (Gamification)", level=2)

    add_heading(doc, "5.7.1 מד התקדמות יומי", level=3)
    add_bullet(doc, "Progress Bar בראש המסך.")
    add_bullet(doc, "יעד יומי הניתן להגדרה ע\u201di המשתמש (ברירת מחדל: 10 ציטוטים ביום).")
    add_bullet(doc, "סופר ציטוטים שעברו ל\u201cרלוונטי\u201d או \u201cלא רלוונטי\u201d היום.")
    add_bullet(doc, "אחוז התקדמות + מספר אבסולוטי.")

    add_heading(doc, "5.7.2 Streak (רצף ימים)", level=3)
    add_bullet(doc, "ספירת רצף ימים עוקבים בהם המשתמש עמד ביעד היומי.")
    add_bullet(doc, "אייקון להבה + מספר הימים מוצג בולט במסך הראשי.")
    add_bullet(doc, "אם הרצף נשבר – הודעה עדינה מעודדת חידוש.")
    add_bullet(doc, "תיעוד היסטוריית streaks.")

    add_heading(doc, "5.7.3 סטטיסטיקות וגרפים", level=3)
    add_para(doc, "דשבורד סטטיסטיקות עם:")
    add_bullet(doc, "גרף ציטוטים-ליום ב-7 וב-30 הימים האחרונים.")
    add_bullet(doc, "פילוח ציטוטים לפי X / Y / I / A (Pie / Bar Chart).")
    add_bullet(doc, "זמן ממוצע מ\u201cבביצוע\u201d ועד החלטה.")
    add_bullet(doc, "סך הכל הושלמו / נותרו, פר תיק וגלובלי.")

    add_heading(doc, "5.7.4 אנימציות וצלילים", level=3)
    add_bullet(doc, "Confetti קצר במסך כשהמשתמש מסיים את היעד היומי.")
    add_bullet(doc, "צליל \u201cping\u201d עדין במעבר כרטיס בין עמודות (ניתן לכבות).")
    add_bullet(doc, "מיקרו-אנימציה (Hover/Drop) על השלמת כרטיס.")

    # 6. Non-functional
    add_heading(doc, "6. דרישות לא-פונקציונליות", level=1)
    add_bullet_with_bold_label(doc, "טכנולוגיה: ", "React + Vite, אחסון ב-localStorage, ללא backend.")
    add_bullet_with_bold_label(doc, "שפה: ", "ממשק כולו בעברית, RTL.")
    add_bullet_with_bold_label(doc, "ביצועים: ", "טעינה ראשונית מתחת ל-1 שנייה; תמיכה בלפחות 5,000 ציטוטים בלוח.")
    add_bullet_with_bold_label(doc, "פרטיות: ", "100% מקומי; אין שליחת נתונים החוצה; אין analytics.")
    add_bullet_with_bold_label(doc, "דפדפנים: ", "Chrome / Safari / Firefox עדכניים בלבד.")
    add_bullet_with_bold_label(doc, "נגישות: ", "ניווט מקלדת בסיסי, ניגודיות תקנית.")

    # 7. UX
    add_heading(doc, "7. עיצוב חוויית משתמש (UX)", level=1)
    add_bullet(doc, "סגנון מינימליסטי ונקי, עם דגש ויזואלי על המספרים והיעדים.")
    add_bullet(doc, "צבע ראשי לכל עמודה: \u201cטרם התחיל\u201d (אפור), \u201cבביצוע\u201d (כחול), \u201cרלוונטי\u201d (ירוק), \u201cלא רלוונטי\u201d (אדום).")
    add_bullet(doc, "ה-Tabs של X/Y/I/A בצבעים שונים זה מזה לזיהוי מהיר.")
    add_bullet(doc, "Dark Mode אופציונלי.")
    add_bullet(doc, "תמיכה בקיצורי מקלדת (לדוגמה: N לציטוט חדש, / לחיפוש).")

    # 8. User flows
    add_heading(doc, "8. תרשימי זרימה (User Flows)", level=1)
    add_heading(doc, "8.1 ייבוא CSV", level=3)
    add_para(doc, "המשתמש בוחר תיק \u2192 לוחץ \u201cייבוא CSV\u201d \u2192 בוחר קובץ \u2192 רואה תצוגה מקדימה \u2192 מאשר \u2192 הציטוטים מופיעים ב\u201cטרם התחיל\u201d של התיק.")
    add_heading(doc, "8.2 בחינת ציטוט", level=3)
    add_para(doc, "המשתמש לוחץ על לינק בכרטיס \u201cטרם התחיל\u201d \u2192 הכרטיס עובר ל\u201cבביצוע\u201d והלינק נפתח \u2192 לאחר הקריאה גורר ל\u201cרלוונטי\u201d \u2192 בוחר X/Y/I/A או גורר ל\u201cלא רלוונטי\u201d.")
    add_heading(doc, "8.3 השלמת יעד יומי", level=3)
    add_para(doc, "כשהמשתמש מגיע ליעד היומי \u2192 confetti במסך + עדכון Streak + הודעה קצרה.")

    # 9. Out of scope
    add_heading(doc, "9. מחוץ להיקף (Out of Scope)", level=1)
    add_bullet(doc, "אין משתמשים מרובים, אין שיתוף, אין הרשאות.")
    add_bullet(doc, "אין סנכרון לענן או בין מכשירים.")
    add_bullet(doc, "אין אינטגרציה אוטומטית ל-EPO / USPTO / Espacenet.")
    add_bullet(doc, "אין גרסת מובייל ייעודית בגרסה הראשונה.")

    # 10. Milestones
    add_heading(doc, "10. תוכנית פיתוח (Milestones)", level=1)
    add_table(
        doc,
        headers=["שלב", "תוכן", "פלט נדרש"],
        rows=[
            ["M1 – שלד", "Vite + React, מבנה תיקיות, מתאם localStorage, לוח אחד עם 4 עמודות וכרטיסים בסיסיים.", "אפליקציה רצה מקומית עם CRUD ידני."],
            ["M2 – Tabs + תיקים", "Tabs של X/Y/I/A בעמודת \u201cרלוונטי\u201d, מסך בחירת תיק, ניווט בין תיקים.", "ריבוי תיקים פעיל."],
            ["M3 – CSV + JSON", "ייבוא CSV עם תצוגה מקדימה וולידציה, ייצוא וייבוא JSON, גיבוי יומי.", "מערכת גיבוי שלמה."],
            ["M4 – מוטיבציה", "Progress Bar יומי, Streak, אנימציות, צלילים.", "אלמנטי משחוק פעילים."],
            ["M5 – סטטיסטיקות", "דשבורד עם גרפים (Recharts).", "מסך נתונים שלם."],
            ["M6 – ליטוש", "Dark Mode, קיצורי מקלדת, ליטוש UX.", "גרסה 1.0 יציבה."],
        ],
    )

    # 11. Appendices
    add_heading(doc, "11. נספחים", level=1)

    add_heading(doc, "11.1 דוגמת CSV", level=2)
    add_table(
        doc,
        headers=["publication_number", "title", "abstract", "link", "notes"],
        rows=[
            ["US1234567B2", "Method for X", "A method for...", "https://patents.google.com/patent/US1234567B2", ""],
            ["EP9876543A1", "System for Y", "A system that...", "https://worldwide.espacenet.com/...", "לבדוק טענה 3"],
        ],
    )

    add_heading(doc, "11.2 מילון מונחים – סיווגי X / Y / I / A", level=2)
    add_bullet_with_bold_label(doc, "X – ", "מסמך פוסל בפני עצמו (Particularly relevant alone).")
    add_bullet_with_bold_label(doc, "Y – ", "מסמך פוסל בשילוב עם מסמך אחר (Particularly relevant in combination).")
    add_bullet_with_bold_label(doc, "I – ", "מסמך מעניין / טכנולוגית רקע (Technological background of particular interest).")
    add_bullet_with_bold_label(doc, "A – ", "מסמך רקע כללי (General technological background).")

    add_heading(doc, "11.3 מבנה JSON (Schema)", level=2)
    add_para(doc, "מבנה הנתונים בייצוא JSON:")
    add_bullet(doc, "version – גרסת הסכמה.")
    add_bullet(doc, "exportedAt – תאריך ושעת הייצוא.")
    add_bullet(doc, "settings – הגדרות משתמש (יעד יומי, צלילים פעילים, dark mode).")
    add_bullet(doc, "streak – נתוני רצף ({current, longest, lastQualifiedDate}).")
    add_bullet(doc, "boards[] – מערך תיקים, כל אחד עם: id, name, createdAt, citations[].")
    add_bullet(doc, "citations[] – מערך כרטיסים, כל אחד עם: id, publicationNumber, title, abstract, link, notes, status, subClassification, createdAt, updatedAt, completedAt.")

    out_path = "/Users/benjaminbodner/Documents/סדנאות/רשות_הפטנטים/claude_code_session2/citations_manager/citations_manager_PRD.docx"
    doc.save(out_path)
    print(f"Saved: {out_path}")


if __name__ == "__main__":
    main()
